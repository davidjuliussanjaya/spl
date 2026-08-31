from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "Lampiran_Source_Code_22410100009.docx"

TITLE_LINES = [
    "RANCANG BANGUN APLIKASI SURVEI PENGGUNA LULUSAN",
    "BERBASIS WEB DENGAN DASHBOARD INTERAKTIF",
    "MENGGUNAKAN METODE PIECES",
]

APPENDICES = [
    (
        "Routing Utama Aplikasi",
        "routes/web.php",
        "Definisi rute aplikasi yang menghubungkan autentikasi, dashboard, pengelolaan survei, pengisian responden, dan laporan.",
    ),
    (
        "Controller Dashboard",
        "app/Http/Controllers/DashboardController.php",
        "Controller yang menerima filter periode, fakultas, dan program studi sebelum meneruskan data terolah ke dashboard.",
    ),
    (
        "Service Perhitungan Dashboard",
        "app/Services/DashboardService.php",
        "Logika utama perhitungan indeks kepuasan, rata-rata kategori, distribusi jawaban, statistik responden, dan data drill-down dashboard.",
    ),
    (
        "Tampilan Dashboard Interaktif",
        "resources/views/admin/dashboard/index.blade.php",
        "Antarmuka dashboard interaktif yang menampilkan kartu indikator, filter dinamis, tabel kepuasan, grafik ApexCharts, modal, dan fitur drill-down.",
    ),
    (
        "Controller Survei",
        "app/Http/Controllers/SurveyController.php",
        "Controller untuk pembuatan survei, verifikasi kode akses, penayangan formulir, pengiriman jawaban, pembaruan, dan pembuatan survei massal.",
    ),
    (
        "Service Pengolahan Survei",
        "app/Services/SurveyService.php",
        "Layanan bisnis untuk membuat survei, menyimpan jawaban secara transaksional, membuat arsip permanen, dan mengelola survei massal.",
    ),
    (
        "Validasi Jawaban Survei",
        "app/Http/Requests/SurveySubmitJawabanRequest.php",
        "Validasi masukan responden sebelum data jawaban, informasi perusahaan, dan identitas pengisi diproses oleh aplikasi.",
    ),
    (
        "Formulir Pengisian Survei",
        "resources/views/fill_page.blade.php",
        "Tampilan publik untuk pengisian profil pengguna lulusan dan instrumen penilaian, termasuk rating, pilihan ganda, dan jawaban esai.",
    ),
    (
        "Service Pengelolaan Pertanyaan",
        "app/Services/PertanyaanService.php",
        "Logika pengelolaan kategori, jenis pertanyaan, pilihan jawaban, nilai skala, serta status aktif instrumen survei.",
    ),
    (
        "Seeder Instrumen Survei Pengguna Lulusan",
        "database/seeders/DraftInstrumenUniversitas2026Seeder.php",
        "Pembentukan instrumen, kategori penilaian, butir pertanyaan, dan bobot jawaban skala empat sebagai dasar pengolahan hasil survei.",
    ),
    (
        "Controller Laporan",
        "app/Http/Controllers/ReportController.php",
        "Controller yang menyiapkan filter laporan, ringkasan survei selesai, daftar arsip, detail arsip, dan proses unduh laporan.",
    ),
    (
        "Ekspor dan Perhitungan Laporan",
        "app/Exports/ReportExport.php",
        "Generator laporan spreadsheet yang menghitung distribusi persentase jawaban, subtotal kategori, total, dan rata-rata per program studi.",
    ),
    (
        "Model Survei",
        "app/Models/Survey.php",
        "Model utama survei beserta relasi lulusan, pengguna lulusan, pertanyaan, instrumen, respon jawaban, dan arsip.",
    ),
    (
        "Model Arsip Survei",
        "app/Models/SurveyArsip.php",
        "Model snapshot permanen hasil survei yang menjadi sumber data stabil untuk dashboard dan pelaporan.",
    ),
    (
        "Model Respon Jawaban",
        "app/Models/ResponJawaban.php",
        "Model detail jawaban responden yang menyimpan relasi soal, pilihan jawaban, nilai, dan teks snapshot.",
    ),
]


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, *, top=50, start=80, bottom=50, end=80) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D7DEE8", size="4") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    settings = {
        "top": ("single", color, size),
        "left": ("single", color, size),
        "bottom": ("single", color, size),
        "right": ("single", color, size),
        "insideH": ("nil", color, "0"),
        "insideV": ("single", color, size),
    }
    for edge, (style, edge_color, edge_size) in settings.items():
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), style)
        element.set(qn("w:color"), edge_color)
        element.set(qn("w:sz"), edge_size)


def add_bottom_border(paragraph, color="222222", size="8") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_left_border(paragraph, color="7385A3", size="18") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), size)
    left.set(qn("w:space"), "10")
    left.set(qn("w:color"), color)
    p_bdr.append(left)


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instruction
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.extend((fld_char_begin, instr_text, fld_char_end))


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Mm(297)
    section.page_height = Mm(210)
    section.top_margin = Mm(15)
    section.bottom_margin = Mm(14)
    section.left_margin = Mm(20)
    section.right_margin = Mm(20)
    section.header_distance = Mm(7)
    section.footer_distance = Mm(7)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(0)

    styles = document.styles
    for style_name in ("Title", "Heading 1", "Heading 2"):
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.color.rgb = RGBColor(20, 25, 35)

    section.header.is_linked_to_previous = False
    header = section.header
    header_p = header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_p.add_run("LAMPIRAN SOURCE CODE UTAMA")
    run.font.name = "Arial"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(90, 98, 112)

    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.paragraph_format.tab_stops.add_tab_stop(Cm(24.8), WD_TAB_ALIGNMENT.RIGHT)
    left = footer_p.add_run("Universitas Dinamika")
    left.font.name = "Arial"
    left.font.size = Pt(8)
    left.font.color.rgb = RGBColor(90, 98, 112)
    footer_p.add_run("\t")
    page_label = footer_p.add_run("Halaman ")
    page_label.font.name = "Arial"
    page_label.font.size = Pt(8)
    add_field(footer_p, "PAGE")
    of_label = footer_p.add_run(" dari ")
    of_label.font.name = "Arial"
    of_label.font.size = Pt(8)
    add_field(footer_p, "NUMPAGES")


def add_spacer(document: Document, points: float) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = Pt(points)


def add_cover(document: Document) -> None:
    add_spacer(document, 42)
    for title_line in TITLE_LINES:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(3)
        run = paragraph.add_run(title_line)
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(16)

    add_spacer(document, 34)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("LAMPIRAN SOURCE CODE UTAMA")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)

    add_spacer(document, 30)
    identity = document.add_table(rows=2, cols=3)
    identity.alignment = WD_TABLE_ALIGNMENT.CENTER
    identity.autofit = False
    widths = (Cm(3.1), Cm(0.7), Cm(8.6))
    data = (("NAMA", ":", "David Julius Sanjaya"), ("NIM", ":", "22410100009"))
    for row, values in zip(identity.rows, data):
        for cell, value, width in zip(row.cells, values, widths):
            cell.width = width
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(value)
            r.font.name = "Times New Roman"
            r.font.size = Pt(11.5)
            if value in {"NAMA", "NIM"}:
                r.bold = True
        prevent_row_split(row)

    tbl_pr = identity._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "nil")
        borders.append(node)
    tbl_pr.append(borders)

    add_spacer(document, 45)
    for line in (
        "PROGRAM STUDI S1 SISTEM INFORMASI",
        "FAKULTAS TEKNOLOGI DAN ELEKTRONIKA",
        "UNIVERSITAS DINAMIKA",
        "2026",
    ):
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(line)
        r.bold = True
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)


def add_appendix_list(document: Document) -> None:
    document.add_page_break()
    heading = document.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.space_after = Pt(16)
    run = heading.add_run("DAFTAR LAMPIRAN KODE SUMBER")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    add_bottom_border(heading)

    table = document.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for index, (title, path, _) in enumerate(APPENDICES, start=1):
        row = table.add_row()
        row.cells[0].width = Cm(21.8)
        row.cells[1].width = Cm(3.2)
        prevent_row_split(row)

        left = row.cells[0]
        p_title = left.paragraphs[0]
        p_title.paragraph_format.space_after = Pt(1)
        title_run = p_title.add_run(f"{index}. {title}")
        title_run.bold = True
        title_run.font.name = "Times New Roman"
        title_run.font.size = Pt(10.5)
        p_path = left.add_paragraph()
        p_path.paragraph_format.left_indent = Cm(0.55)
        p_path.paragraph_format.space_after = Pt(6)
        path_run = p_path.add_run(path)
        path_run.font.name = "Consolas"
        path_run.font.size = Pt(8)
        path_run.font.color.rgb = RGBColor(105, 105, 112)

        right = row.cells[1]
        p_ref = right.paragraphs[0]
        p_ref.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        ref_run = p_ref.add_run(f"[Lampiran {index}]")
        ref_run.bold = True
        ref_run.font.name = "Times New Roman"
        ref_run.font.size = Pt(10)

    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "nil")
        borders.append(node)
    tbl_pr.append(borders)


def add_code_table(document: Document, code: str) -> None:
    lines = code.splitlines()
    if code.endswith("\n"):
        lines.append("")

    table = document.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False
    set_table_borders(table)

    for line_number, raw_line in enumerate(lines, start=1):
        row = table.add_row()
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        prevent_row_split(row)
        number_cell, code_cell = row.cells
        number_cell.width = Cm(1.2)
        code_cell.width = Cm(23.8)
        number_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        code_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(number_cell, "F2F5F9")
        set_cell_shading(code_cell, "FBFCFE")
        set_cell_margins(number_cell, top=12, start=55, bottom=12, end=70)
        set_cell_margins(code_cell, top=12, start=90, bottom=12, end=60)

        number_p = number_cell.paragraphs[0]
        number_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        number_p.paragraph_format.space_after = Pt(0)
        number_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        number_p.paragraph_format.line_spacing = Pt(8.4)
        number_run = number_p.add_run(str(line_number))
        number_run.font.name = "Consolas"
        number_run.font.size = Pt(7)
        number_run.font.color.rgb = RGBColor(115, 124, 138)

        code_p = code_cell.paragraphs[0]
        code_p.paragraph_format.space_after = Pt(0)
        code_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        code_p.paragraph_format.line_spacing = Pt(8.4)
        code_run = code_p.add_run(raw_line.expandtabs(4) if raw_line else " ")
        code_run.font.name = "Consolas"
        code_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        code_run.font.size = Pt(7)
        code_run.font.color.rgb = RGBColor(25, 29, 36)


def add_appendix(document: Document, index: int, title: str, relative_path: str, description: str) -> None:
    document.add_page_break()

    heading = document.add_paragraph()
    heading.paragraph_format.space_after = Pt(2)
    run = heading.add_run(f"{index}. {title}")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(13)

    path_paragraph = document.add_paragraph()
    path_paragraph.paragraph_format.space_after = Pt(8)
    path_run = path_paragraph.add_run(f"File path: {relative_path}")
    path_run.font.name = "Consolas"
    path_run.font.size = Pt(8.5)
    path_run.font.color.rgb = RGBColor(80, 85, 94)
    add_bottom_border(path_paragraph, color="444444", size="6")

    description_paragraph = document.add_paragraph()
    description_paragraph.paragraph_format.left_indent = Cm(0.35)
    description_paragraph.paragraph_format.right_indent = Cm(0.35)
    description_paragraph.paragraph_format.space_before = Pt(5)
    description_paragraph.paragraph_format.space_after = Pt(10)
    description_run = description_paragraph.add_run(description)
    description_run.italic = True
    description_run.font.name = "Times New Roman"
    description_run.font.size = Pt(10)
    description_run.font.color.rgb = RGBColor(70, 77, 90)
    add_left_border(description_paragraph)

    source_path = ROOT / relative_path
    code = source_path.read_text(encoding="utf-8", errors="replace")
    add_code_table(document, code)


def set_update_fields_on_open(document: Document) -> None:
    settings = document.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def main() -> None:
    missing = [path for _, path, _ in APPENDICES if not (ROOT / path).is_file()]
    if missing:
        raise FileNotFoundError(f"Source files not found: {', '.join(missing)}")

    document = Document()
    configure_document(document)
    document.core_properties.title = "Lampiran Source Code Utama - David Julius Sanjaya"
    document.core_properties.author = "David Julius Sanjaya"
    document.core_properties.subject = "Rancang Bangun Aplikasi Survei Pengguna Lulusan Berbasis Web dengan Dashboard Interaktif Menggunakan Metode PIECES"
    document.core_properties.keywords = "source code, survei pengguna lulusan, dashboard, PIECES, Laravel"

    add_cover(document)
    add_appendix_list(document)
    for index, appendix in enumerate(APPENDICES, start=1):
        add_appendix(document, index, *appendix)

    set_update_fields_on_open(document)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
